from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates

import gspread
from google.oauth2.service_account import Credentials

import uuid
from docx import Document

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# -----------------------------
# GOOGLE SHEETS CONFIG
# -----------------------------
SCOPES = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
SHEET_NAME = "Table Tent Menu DB"

menu_db = None  # lazy-loaded cache


# -----------------------------
# GOOGLE SHEETS CONNECTION
# -----------------------------
def get_sheet():
    creds = Credentials.from_service_account_file(
        "service_account.json",
        scopes=SCOPES
    )

    client = gspread.authorize(creds)
    sheet = client.open(SHEET_NAME).sheet1
    return sheet


def load_menu():
    sheet = get_sheet()
    records = sheet.get_all_records()

    menu = {}
    for r in records:
        menu[r["item_name"].lower()] = r

    return menu


def get_menu():
    global menu_db
    if menu_db is None:
        menu_db = load_menu()
    return menu_db


# -----------------------------
# WEB UI
# -----------------------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# -----------------------------
# CORE GENERATION LOGIC
# -----------------------------
def generate_table_tents(items):
    menu = get_menu()

    output_file = f"table_tents_{uuid.uuid4().hex}.docx"
    doc = Document()

    for item in items:
        record = menu.get(item.lower())

        if not record:
            continue

        doc.add_paragraph(record["display_name"])
        doc.add_paragraph(record["description"])
        doc.add_paragraph(f"Contains: {record['allergens']}")
        doc.add_paragraph(record["tags"])
        doc.add_page_break()

    doc.save(output_file)
    return output_file


# -----------------------------
# TEMP TEST DATA (replace later with AI)
# -----------------------------
def parse_order():
    return ["grilled chicken", "caesar salad"]


# -----------------------------
# API ENDPOINT
# -----------------------------
@app.get("/generate")
def generate():
    items = parse_order()
    file_path = generate_table_tents(items)

    return FileResponse(
        file_path,
        filename="table_tents.docx"
    )
